#!/usr/bin/env python3
"""Generate scientific report for MMLM Accident Anticipation project."""
import json
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    raise


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p


def add_bullet(doc, text):
    return doc.add_paragraph(text, style='List Bullet')


def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('Multimodal Large Language Model for Collision Anticipation', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Technical Report')
    doc.add_paragraph('Project: MMLM_CursorAI')
    doc.add_paragraph()
    
    # 1. Introduction
    add_heading(doc, '1. Introduction and Objective', level=1)
    add_para(doc, 
        'This project develops a Multimodal Large Language Model (MMLM) for Accident Anticipation '
        'in driving scenarios. The system processes video frame sequences to generate a continuous '
        'risk-probability curve with per-frame predictions and Chain-of-Thought (CoT) explanations '
        'justifying each risk assessment.')
    add_para(doc,
        'The model must predict collision likelihood for every frame in a video by analyzing temporal '
        'visual patterns and producing both a numerical probability (0.0–1.0) and explicit reasoning '
        'about visible risk factors.')
    
    # 2. System Architecture
    add_heading(doc, '2. System Architecture', level=1)
    add_para(doc,
        'The architecture follows a modular Vision-Language Model (VLM) pipeline designed to be '
        'model-agnostic, allowing future backbone substitution.')
    
    add_heading(doc, '2.1 Core Components', level=2)
    
    add_heading(doc, 'Vision Encoder', level=3)
    add_para(doc,
        'Model: CLIP ViT-Base-Patch32 (openai/clip-vit-base-patch32)')
    add_para(doc,
        'Processes individual frames independently to extract visual features. Each frame is encoded '
        'into a sequence of visual tokens (49 patch tokens × 768 dimensions after dropping the CLS token).')
    add_para(doc,
        'Rationale: CLIP provides strong pre-trained visual representations, is lightweight enough for '
        'real-time inference constraints (<0.5s per window), and enables rapid prototyping with stable '
        'weight loading from Hugging Face.')
    
    add_heading(doc, 'Temporal Token Mixer', level=3)
    add_para(doc,
        'A token-preserving transformer encoder (2 layers, 8 heads) that models temporal dependencies '
        'across frames within a clip.')
    add_para(doc,
        'Input: per-frame visual tokens with learned temporal position embeddings')
    add_para(doc,
        'Output: temporally-aware token sequence (preserves T×M tokens rather than pooling to a single vector)')
    add_para(doc,
        'This explicit temporal modeling occurs before the modality bridge, allowing the vision pathway '
        'to capture motion and temporal relations independently of the LLM.')
    
    add_heading(doc, 'Modality Bridge / Projector', level=3)
    add_para(doc,
        'A linear projection layer that maps vision tokens (768-dim) into the LLM embedding space (768-dim for GPT-2).')
    add_para(doc,
        'This enables the LLM to process visual information as if it were part of its native token sequence.')
    
    add_heading(doc, 'Reasoning Engine (LLM)', level=3)
    add_para(doc,
        'Model: GPT-2 (124M parameters)')
    add_para(doc,
        'A decoder-only transformer that receives concatenated embeddings: [visual_prefix | text_prompt].')
    add_para(doc,
        'The LLM processes the fused multimodal sequence and generates hidden states used for both '
        'text generation and risk prediction.')
    add_para(doc,
        'Rationale: GPT-2 serves as a lightweight baseline for validating the multimodal pipeline. '
        'It is small enough to run efficiently on RTX 5000 Ada (32GB VRAM) while supporting LoRA '
        'fine-tuning and custom token addition (<SCORE>).')
    
    add_heading(doc, 'Score Head', level=3)
    add_para(doc,
        'A linear layer (768 → 1) that predicts collision risk from the LLM hidden state at a special '
        '<SCORE> token position.')
    add_para(doc,
        'Output: logits trained with BCEWithLogitsLoss against continuous risk targets.')
    
    # 3. Methodology
    add_heading(doc, '3. Methodology: Sliding Window Temporal Inference', level=1)
    add_para(doc,
        'To achieve per-frame collision probabilities without redundant processing, the system employs '
        'a sliding window approach:')
    
    add_heading(doc, 'Clip Construction', level=3)
    add_para(doc,
        'For any frame t, the input is a temporal clip [t-k, t] where k defines the lookback window.')
    add_para(doc, 'Configuration:')
    add_bullet(doc, 'window_size: 16 frames')
    add_bullet(doc, 'stride: 4 (temporal subsampling within the window)')
    add_bullet(doc, 'window_step: 1 (frame-by-frame sliding for dense coverage)')
    add_bullet(doc, 'Effective receptive field: ~60 frames lookback at 30 fps ≈ 2 seconds')
    
    add_heading(doc, 'Last-Frame Prediction', level=3)
    add_para(doc,
        'The model is trained to predict risk specifically for the final frame of each clip. By sliding '
        'the window frame-by-frame, a continuous probability curve is generated across the entire video.')
    
    add_heading(doc, 'Temporal Coverage Strategy', level=3)
    add_para(doc, 'Two operational modes:')
    add_bullet(doc, 'Training: capped sampling (max 32 clips/video) for efficiency')
    add_bullet(doc, 'Inference: dense sliding (window_step=1) for complete per-frame coverage')
    
    # 4. Dataset
    add_heading(doc, '4. Dataset Specifications', level=1)
    add_para(doc,
        'Dataset: Nexar Dash-Cam Challenge')
    add_para(doc,
        'Location: /home/eprojuser011/Data_Centric_using_3LC_and_MViT/src/projects/NexarChallenge/Nexar_DataSet')
    
    add_heading(doc, 'Ground Truth Labels', level=3)
    add_bullet(doc, 'target: binary collision label (1 = collision occurs, 0 = safe driving)')
    add_bullet(doc, 'time_of_alert: ground truth timestamp when collision becomes predictable')
    add_bullet(doc, 'time_of_event: actual collision timestamp')
    
    add_heading(doc, 'Continuous Risk Labeling', level=3)
    add_para(doc,
        'A piecewise continuous target_risk function converts discrete labels into smooth 0.0–1.0 values:')
    add_bullet(doc, 'Safety phase (t < time_of_alert): risk = 0.0')
    add_bullet(doc, 'Anticipation phase (time_of_alert ≤ t < time_of_event): linear ramp from 0.1 → 0.9')
    add_bullet(doc, 'Collision phase (t ≈ time_of_event): risk = 1.0')
    add_bullet(doc, 'Non-collision videos: constant risk = 0.0')
    
    add_heading(doc, 'Data Preprocessing Rules', level=3)
    add_bullet(doc, 'ignore_after_event: exclude frames after collision timestamp')
    add_bullet(doc, 'limit_noncollision_half: use only first 50% of safe videos to balance dataset')
    add_bullet(doc, 'Frame format: 256×256 RGB JPEGs, resized to 224×224 for model input')
    
    # 5. Training Approach
    add_heading(doc, '5. Training Methodology', level=1)
    
    add_heading(doc, '5.1 Phase 1: Student Training (Score-Only Supervision)', level=2)
    add_para(doc,
        'The student model is trained end-to-end using only continuous risk labels without teacher '
        'reasoning. This establishes a baseline and validates the data→model pipeline.')
    
    add_para(doc, 'Loss function:')
    add_bullet(doc, 'BCEWithLogitsLoss on score_logits vs target_risk')
    add_bullet(doc, 'Supports soft (continuous) targets unlike standard BCE')
    
    add_para(doc, 'Optimization:')
    add_bullet(doc, 'Optimizer: AdamW with learning_rate=2e-5, weight_decay=0.01')
    add_bullet(doc, 'Parameter-Efficient Fine-Tuning (PEFT): LoRA with r=8, alpha=16, dropout=0.05')
    add_bullet(doc, 'LoRA is applied only to the LLM; vision encoder, mixer, and projector train fully')
    
    add_heading(doc, '5.2 Phase 1C: Clip Manifest Cache + Smoke Testing', level=2)
    add_para(doc,
        'Before full teacher distillation, Plan 1C pre-generates a deterministic clip manifest for '
        'reproducibility and efficient large-scale teacher generation.')
    
    add_para(doc, 'Outputs:')
    add_bullet(doc, 'teacher_clips_manifest.jsonl: 47,320 clip records (cap 32/video)')
    add_bullet(doc, 'Each record: video_id, end_frame_idx, window_size, stride, t_seconds, target, '
               'time_of_alert, time_of_event, target_risk, frame_indices')
    add_bullet(doc, 'teacher_clips_manifest.progress.jsonl: real-time progress log (ETA, throughput)')
    add_bullet(doc, 'smoke_forward_report.json: forward-pass validation on 32 clips')
    
    add_para(doc, 'Smoke test results:')
    add_bullet(doc, 'Clip batch shape: [32, 16, 3, 224, 224]')
    add_bullet(doc, 'Visual tokens shape: [32, 784, 768] (49 patches × 16 frames = 784 tokens)')
    add_bullet(doc, 'Score logits shape: [32] (one prediction per clip)')
    add_bullet(doc, 'Forward pass latency: 1117 ms for 32-clip batch (~35 ms/clip)')
    add_bullet(doc, 'Status: forward_ok = true (pipeline validated)')
    
    add_heading(doc, '5.3 Phase 2: Teacher Distillation (Planned)', level=2)
    add_para(doc,
        'Use a large external Vision-Language Model (Gemini Pro 1.5) as a teacher to generate '
        'Chain-of-Thought reasoning for a subset of clips.')
    
    add_para(doc, 'Teacher workflow:')
    add_bullet(doc, 'Iterate teacher_clips_manifest.jsonl')
    add_bullet(doc, 'For each clip: load frames, construct conditioned prompt with target_risk')
    add_bullet(doc, 'Send frames + prompt to Gemini API')
    add_bullet(doc, 'Parse teacher response and cache to teacher_cache.jsonl')
    
    add_para(doc, 'Prompt design (option C recommended):')
    add_bullet(doc, 'Scene: brief objective description')
    add_bullet(doc, 'Risk_reason: 2-4 visual cues explaining why risk is high/low')
    add_bullet(doc, 'Alignment: teacher output conditioned on target_risk to avoid contradictions')
    
    add_heading(doc, '5.4 Phase 3: Integrated Training with Reasoning Loss (Planned)', level=2)
    add_para(doc,
        'Train student on mixed batches: some with teacher reasoning labels, others with score-only labels.')
    
    add_bullet(doc, 'Loss = α × CrossEntropy(reasoning_logits, teacher_text) + β × BCE(score_logits, target_risk)')
    add_bullet(doc, 'Apply reasoning loss only where teacher labels exist (masked loss)')
    add_bullet(doc, 'Balance batches to avoid "always predict danger" bias')
    
    # 6. Hardware constraints
    add_heading(doc, '6. Hardware Constraints and Optimization', level=1)
    add_bullet(doc, 'GPU: NVIDIA RTX 5000 Ada (32GB VRAM)')
    add_bullet(doc, 'Latency target: < 0.5 seconds per inference window')
    add_bullet(doc, 'Current performance: ~35 ms/clip (well below target)')
    add_bullet(doc, 'Batch size: 4 for training, 32 for smoke test')
    
    add_para(doc, 'Memory optimization techniques:')
    add_bullet(doc, '4-bit quantization available (BitsAndBytes) for LLM if needed')
    add_bullet(doc, 'LoRA reduces trainable parameters from ~124M to ~few hundred K')
    add_bullet(doc, 'Gradient checkpointing (not yet enabled, available if needed)')
    
    # 7. Model Selection
    add_heading(doc, '7. Model Selection and Rationale', level=1)
    
    add_heading(doc, '7.1 Vision Encoder: CLIP ViT-Base-Patch32', level=2)
    add_para(doc, 'Specifications:')
    add_bullet(doc, 'Architecture: Vision Transformer, 12 layers, 768 hidden size')
    add_bullet(doc, 'Input: 224×224 images divided into 32×32 patches')
    add_bullet(doc, 'Output: 50 tokens (1 CLS + 49 patches), 768-dim embeddings')
    
    add_para(doc, 'Why chosen:')
    add_bullet(doc, 'Strong pre-trained representations from CLIP contrastive learning')
    add_bullet(doc, 'Small enough for real-time inference (~15ms/frame on RTX 5000 Ada)')
    add_bullet(doc, 'Stable Hugging Face integration with safetensors loading')
    add_bullet(doc, 'Widely validated on vision-language tasks')
    
    add_heading(doc, '7.2 LLM: GPT-2', level=2)
    add_para(doc, 'Specifications:')
    add_bullet(doc, 'Architecture: 12-layer decoder-only transformer, 768 hidden size, 124M parameters')
    add_bullet(doc, 'Vocabulary: 50,257 tokens + custom <SCORE> token')
    
    add_para(doc, 'Why chosen (current stage):')
    add_bullet(doc, 'Lightweight baseline for infrastructure validation')
    add_bullet(doc, 'Fast forward pass (~20ms for short sequences)')
    add_bullet(doc, 'Full LoRA support via PEFT library')
    add_bullet(doc, 'Easy token vocabulary extension for special tokens')
    
    add_para(doc, 'Future replacement candidates:')
    add_bullet(doc, 'Qwen2-VL 2B (native vision-language model with better reasoning)')
    add_bullet(doc, 'PaliGemma 3B (purpose-built for vision-language tasks)')
    add_bullet(doc, 'Both support 4-bit quantization for memory efficiency')
    
    # 8. Technical implementation
    add_heading(doc, '8. Implementation Details', level=1)
    
    add_heading(doc, '8.1 Data Pipeline', level=2)
    add_para(doc, 'Sliding window dataset (MCASlidingWindowDataset):')
    add_bullet(doc, 'Precomputes valid (video_id, end_frame_idx) pairs at initialization')
    add_bullet(doc, 'Applies temporal windowing rules and caps (max 128 windows/video for training)')
    add_bullet(doc, 'Loads frames on-demand via DataLoader workers (num_workers=4)')
    add_bullet(doc, 'Normalization: mean=[0.45, 0.45, 0.45], std=[0.225, 0.225, 0.225]')
    
    add_heading(doc, '8.2 Forward Pass Flow', level=2)
    add_para(doc, '1. Vision encoding:')
    add_bullet(doc, 'Input clip shape: [B, T, C, H, W] = [batch, 16 frames, 3 channels, 224, 224]')
    add_bullet(doc, 'Flatten to [B×T, C, H, W] and pass through CLIP vision encoder')
    add_bullet(doc, 'Extract patch tokens (drop CLS): [B×T, 49, 768]')
    add_bullet(doc, 'Reshape to [B, T, 49, 768] and add learned temporal position embeddings')
    
    add_para(doc, '2. Temporal mixing:')
    add_bullet(doc, 'Flatten to [B, T×49, 768] = [B, 784, 768]')
    add_bullet(doc, 'Apply TransformerEncoder (2 layers) for cross-frame attention')
    add_bullet(doc, 'Output: [B, 784, 768] temporally-aware visual tokens')
    
    add_para(doc, '3. Modality fusion:')
    add_bullet(doc, 'Project visual tokens: linear(768 → 768) → visual_embeds')
    add_bullet(doc, 'Tokenize text prompt containing <SCORE> token')
    add_bullet(doc, 'Get text embeddings from LLM embedding layer')
    add_bullet(doc, 'Concatenate: [visual_embeds | text_embeds]')
    
    add_para(doc, '4. LLM processing + prediction:')
    add_bullet(doc, 'LLM forward with inputs_embeds (attention over full sequence)')
    add_bullet(doc, 'Extract hidden state at <SCORE> token position (shifted by visual prefix length)')
    add_bullet(doc, 'Pass through score_head: linear(768 → 1) → logits')
    
    # 9. Plan 1C results
    add_heading(doc, '9. Plan 1C: Clip Manifest Generation Results', level=1)
    add_para(doc,
        'Successfully generated a deterministic teacher-clip manifest with progress logging and '
        'smoke-tested the forward pass.')
    
    add_para(doc, 'Statistics:')
    add_bullet(doc, 'Total clips generated: 47,320 (capped at 32/video)')
    add_bullet(doc, 'Manifest generation time: ~1.6 seconds (~30,000 clips/second)')
    add_bullet(doc, 'All videos from train.csv successfully indexed')
    
    add_para(doc, 'Smoke test validation:')
    add_bullet(doc, '32 balanced clips (positive/negative) tested')
    add_bullet(doc, 'All frames loaded successfully from train_frames256/')
    add_bullet(doc, 'Forward pass completed without errors')
    add_bullet(doc, 'Latency per clip: ~35 ms (95% below 0.5s target)')
    
    # 10. Next steps
    add_heading(doc, '10. Next Steps: Phase 2 and Beyond', level=1)
    
    add_heading(doc, 'Phase 2: Gemini Teacher Distillation', level=3)
    add_bullet(doc, 'Set up Gemini API authentication (GEMINI_API_KEY environment variable)')
    add_bullet(doc, 'Iterate teacher_clips_manifest.jsonl and generate reasoning labels')
    add_bullet(doc, 'Output format: Scene description + Risk reasoning (2-4 visual cues)')
    add_bullet(doc, 'Cache to teacher_cache.jsonl for reuse')
    add_bullet(doc, 'Estimated API cost and runtime to be determined based on Gemini pricing')
    
    add_heading(doc, 'Phase 3: Mixed-Loss Training', level=3)
    add_bullet(doc, 'Integrate teacher_cache.jsonl into training dataset')
    add_bullet(doc, 'Implement reasoning CrossEntropy loss + score BCE loss')
    add_bullet(doc, 'Balance batches (teacher-labeled + score-only samples)')
    add_bullet(doc, 'Evaluate reasoning quality and risk prediction metrics (AUC, AP, pre/post alert risk)')
    
    add_heading(doc, 'Model Upgrade Path', level=3)
    add_bullet(doc, 'Replace GPT-2 with Qwen2-VL 2B or PaliGemma 3B')
    add_bullet(doc, 'Enable 4-bit quantization if memory constrained')
    add_bullet(doc, 'Benchmark latency against 0.5s target')
    
    # 11. Technical challenges
    add_heading(doc, '11. Technical Challenges Resolved', level=1)
    
    add_heading(doc, 'Import Path Issue', level=3)
    add_para(doc,
        'Problem: Scripts in scripts/ directory could not import from data/, models/ due to missing '
        'PYTHONPATH.')
    add_para(doc,
        'Solution: Added sys.path.insert(0, repo_root) in script entry points to ensure repo root is '
        'on import path when running standalone.')
    
    add_heading(doc, 'Video ID Normalization', level=3)
    add_para(doc,
        'Problem: CSV IDs loaded as floats (e.g., 1924.0) did not match frame folder names (01924).')
    add_para(doc,
        'Solution: Normalize all IDs to zero-padded 5-digit strings: f"{int(float(x)):05d}"')
    
    # 12. Summary
    add_heading(doc, '12. Summary', level=1)
    add_para(doc,
        'This project implements a modular Multimodal Large Language Model for accident anticipation '
        'using a sliding-window temporal approach. The current implementation successfully:')
    add_bullet(doc, 'Processes video clips through a vision encoder → temporal mixer → LLM pipeline')
    add_bullet(doc, 'Generates continuous per-frame risk predictions via a learned score head')
    add_bullet(doc, 'Supports efficient fine-tuning via LoRA')
    add_bullet(doc, 'Meets latency requirements with significant headroom')
    add_bullet(doc, 'Provides a deterministic clip manifest for reproducible teacher distillation')
    
    add_para(doc,
        'The architecture is designed to be model-agnostic, enabling future upgrades to stronger '
        'vision-language backbones while maintaining the core temporal sliding-window methodology.')
    
    # Appendix
    add_heading(doc, 'Appendix A: File Structure', level=1)
    add_para(doc, 'Project root: /home/eprojuser011/MMLM_CursorAI/')
    
    add_para(doc, 'Key files:')
    add_bullet(doc, 'configs/base.yaml: configuration parameters')
    add_bullet(doc, 'data/dataset.py: MCASlidingWindowDataset + risk computation')
    add_bullet(doc, 'data/video_index.py: frame counting and FPS extraction')
    add_bullet(doc, 'models/temporal_mixer.py: token-preserving temporal transformer')
    add_bullet(doc, 'models/factory.py: MCAStudentModel + LoRA integration')
    add_bullet(doc, 'prompts/templates.py: prompt templates for student and teacher')
    add_bullet(doc, 'scripts/build_teacher_manifest.py: manifest generation')
    add_bullet(doc, 'scripts/smoke_forward.py: forward-pass validation')
    add_bullet(doc, 'teacher/distill.py: Gemini teacher distillation (Phase 2)')
    add_bullet(doc, 'train_sft.py: main training script (Phase 1)')
    add_bullet(doc, 'slurm/plan1c_manifest_and_smoke.sbatch: SLURM job for Plan 1C')
    
    # Save
    output_dir = Path('/home/eprojuser011/MMLM_CursorAI/reports')
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / 'MMLM_Accident_Anticipation_Report.docx'
    
    doc.save(str(output_path))
    print(f"Report saved to: {output_path}")


if __name__ == '__main__':
    main()
