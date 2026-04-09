"""
Generate progress report: MMLM-Based Collision Anticipation
Prompt Engineering & Teacher Dataset Construction

Run:
    py -3 -m scripts.generate_progress_report
Output:
    MMLM_CursorAI/outputs/progress_report.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def _heading(doc, text, level, space_before=12, space_after=4):
    p = doc.add_heading(text, level=level)
    fmt = p.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    return p


def _para(doc, text="", bold_parts=None, space_after=6, indent=None):
    """Add a paragraph. bold_parts is a list of substrings to embolden."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    if not bold_parts:
        run = p.add_run(text)
        _set_font(run)
    else:
        remaining = text
        for bp in bold_parts:
            idx = remaining.find(bp)
            if idx == -1:
                continue
            if idx > 0:
                r = p.add_run(remaining[:idx])
                _set_font(r)
            r = p.add_run(bp)
            _set_font(r, bold=True)
            remaining = remaining[idx + len(bp):]
        if remaining:
            r = p.add_run(remaining)
            _set_font(r)
    return p


def _bullet(doc, text, bold_parts=None, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    if not bold_parts:
        run = p.add_run(text)
        _set_font(run)
    else:
        remaining = text
        for bp in bold_parts:
            idx = remaining.find(bp)
            if idx == -1:
                continue
            if idx > 0:
                r = p.add_run(remaining[:idx])
                _set_font(r)
            r = p.add_run(bp)
            _set_font(r, bold=True)
            remaining = remaining[idx + len(bp):]
        if remaining:
            r = p.add_run(remaining)
            _set_font(r)
    return p


def _shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _cell_text(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"


def _add_experiment_table(doc):
    """Add the main experiment results table."""
    headers = ["Run", "Prompt", "TP", "FP", "TN", "FN", "Acc", "Notes"]
    rows = [
        ("v1",  "PROMPT_F (oracle-aligned, GPT-4o)", "-", "-", "-", "-", "N/A",   "Pipeline validation; GT injected"),
        ("v2",  "PROMPT_G (6-step CoT)",              "5", "3", "6", "4", "61%",   "Best blind baseline"),
        ("v3",  "PROMPT_H (+lane discipline, 8-step)","2", "1", "8", "6", "56%",   "Too conservative"),
        ("v3.1a","H + priority rules override",       "8", "7", "2", "1", "56%",   "Too aggressive"),
        ("v3.1d","H + RAPID closing threshold",       "5", "4", "5", "4", "56%",   "Better balance, < v2"),
        ("v4",  "PROMPT_I (enhanced PROMPT_G)",       "6", "8", "1", "2", "39%",   "Worst; hyper-sensitive to merges"),
        ("v5",  "PROMPT_J (reordered 8-stage)",       "-", "-", "-", "-", "~33%",  "Partial run (9/18 clips)"),
        ("v6",  "PROMPT_K (4-stage PERCEIVE/TRACK…)", "5", "4", "5", "4", "56%",   "Re-balanced, still < v2"),
        ("v7",  "ORACLE_DEEP_V3 (GT-injected)",       "-", "-", "-", "-", "N/A",   "Not blind; forensic audit for training data"),
        ("v8",  "Semi-Oracle Debate (PROMPT_G + Debate)", "-","-","-","-","TBD",  "In progress"),
    ]

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Column widths
    widths = [0.5, 2.3, 0.35, 0.35, 0.35, 0.35, 0.45, 2.1]
    for i, col in enumerate(table.columns):
        for cell in col.cells:
            cell.width = Inches(widths[i])

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        _cell_text(hdr_row.cells[i], h, bold=True, size=9)
        _shade_cell(hdr_row.cells[i], "4472C4")
        hdr_row.cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        is_best = row_data[0] == "v2"
        is_inprogress = row_data[0] == "v8"
        for c_idx, val in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if c_idx not in (1, 7) else WD_ALIGN_PARAGRAPH.LEFT
            _cell_text(row.cells[c_idx], val, bold=is_best, size=9, align=align)
            if is_best:
                _shade_cell(row.cells[c_idx], "E2EFDA")   # light green for best
            elif is_inprogress:
                _shade_cell(row.cells[c_idx], "FFF2CC")   # light yellow for in-progress
            elif r_idx % 2 == 0:
                _shade_cell(row.cells[c_idx], "F2F2F2")   # alternating grey

    doc.add_paragraph()  # spacing after table


# ---------------------------------------------------------------------------
# Main document builder
# ---------------------------------------------------------------------------

def build_report(output_path: Path):
    doc = Document()

    # Page margins (narrow)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)

    # -----------------------------------------------------------------------
    # TITLE
    # -----------------------------------------------------------------------
    title = doc.add_heading("Progress Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("MMLM-Based Collision Anticipation: Prompt Engineering & Teacher Dataset Construction")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(16)
    for run in subtitle.runs:
        _set_font(run, size=12, italic=True)

    # -----------------------------------------------------------------------
    # 1. INTRODUCTION
    # -----------------------------------------------------------------------
    _heading(doc, "1. Introduction", 1)

    _para(
        doc,
        "This report summarizes the prompt engineering phase of an ongoing thesis project aimed at building "
        "a teacher-student architecture for ego-vehicle collision anticipation from dashcam footage. "
        "The system uses a large Multimodal Language Model (MMLM) -- specifically Gemini-3 Pro, accessed via "
        "OpenRouter -- as a teacher that generates structured Chain-of-Thought (CoT) reasoning over 16-frame "
        "dashcam clips (approximately 2 seconds of video, sampled at 8 fps). "
        "The teacher outputs a binary collision verdict (YES/NO) together with detailed, structured reasoning "
        "across six analytical dimensions: scene context, dynamic objects, temporal motion analysis, "
        "occlusion check, time-to-contact estimate, and verdict reasoning. "
        "This output will later serve as a distillation signal to train a lightweight student model capable "
        "of real-time deployment without requiring a frontier LLM at inference time.",
        bold_parts=["Gemini-3 Pro", "Chain-of-Thought (CoT)", "YES/NO"],
    )

    _para(
        doc,
        "All experiments in this phase were conducted on a fixed 18-clip test set drawn from the Nexar "
        "real-world dashcam dataset: 9 True Positive (TP) clips sampled at three time-to-event horizons "
        "(0.5 s, 1.0 s, 1.5 s before the collision), and 9 True Negative (TN) clips sampled from the "
        "midpoint of non-collision videos. Each clip was presented as 16 JPEG frames at 256 px resolution. "
        "The primary metric is balanced accuracy -- equal weight on TP recall and TN specificity -- "
        "reflecting the safety-critical requirement of minimising both false alarms and missed collisions.",
        bold_parts=["18-clip test set", "Nexar", "balanced accuracy"],
    )

    # -----------------------------------------------------------------------
    # 2. PROMPT EVOLUTION AND EXPERIMENT RESULTS
    # -----------------------------------------------------------------------
    _heading(doc, "2. Prompt Evolution and Experiment Results", 1)

    _para(
        doc,
        "We ran ten distinct experimental configurations over the same 18 clips, progressively refining "
        "the teacher prompt design. Each version introduced targeted changes motivated by systematic "
        "review of previous errors. The key results are summarised in Table 1 below.",
    )

    doc.add_paragraph()
    caption = doc.add_paragraph("Table 1: Experiment results across all prompt versions (18-clip test set, Nexar).")
    caption.paragraph_format.space_after = Pt(4)
    for run in caption.runs:
        _set_font(run, size=9, italic=True)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_experiment_table(doc)

    _para(
        doc,
        "v1 (PROMPT_F, GPT-4o, oracle-injected): The initial baseline injected the ground-truth label "
        "into the prompt, making it an oracle experiment used purely for pipeline validation. It established "
        "the JSONL/XLSX output pipeline and JSON parsing logic.",
        bold_parts=["v1"],
    )

    _para(
        doc,
        "v2 (PROMPT_G, blind, Gemini-3 Pro): Designed from scratch using the Gemini Prompting Guide 101 "
        "principles (Persona + Task + Context + Format + Constraints) with a 6-step guided CoT structure. "
        "The model received no ground-truth information. This run produced the best balanced result: "
        "TP=5, FP=3, TN=6, FN=4, Accuracy ~61%. The output schema -- scene_context, dynamic_objects, "
        "temporal_analysis, occlusion_check, time_to_contact, collision_verdict, confidence, "
        "verdict_reasoning -- became the uniform reference format for all subsequent versions.",
        bold_parts=["v2", "Accuracy ~61%", "best balanced result"],
    )

    _para(
        doc,
        "v3 (PROMPT_H, 8-step CoT): Extended PROMPT_G with four additional steps motivated by human review "
        "of v2 errors: ego trajectory analysis, three-phase temporal analysis (early / mid / late frames), "
        "lane discipline check, and an explicit collision-target field. However, the lane discipline step "
        "introduced a systematic bias -- the model used 'both vehicles are centered in their lane' as "
        "sufficient evidence for NO, causing TP recall to collapse (TP=2, Accuracy 56%).",
        bold_parts=["v3", "lane discipline step"],
    )

    _para(
        doc,
        "v3.1a-e (PROMPT_H patched, five iterations): Five successive attempts to recover TP recall while "
        "preserving TN specificity by adding priority rules, qualifying closing-speed thresholds, and "
        "adjusting inference parameters (temperature=0, frame_size=512). Every change shifted the YES/NO "
        "decision boundary without achieving balanced improvement; accuracies oscillated between 44% and 56%. "
        "The fundamental finding was that textual threshold descriptions cannot give the model reliably "
        "better visual discrimination at 256px resolution.",
        bold_parts=["v3.1a-e"],
    )

    _para(
        doc,
        "v4 (PROMPT_I, enhanced PROMPT_G): Folded lane position and three-phase temporal analysis back into "
        "the 6-step PROMPT_G structure. A same-lane rear-end warning injected at the observation step "
        "primed the model to interpret any lane change as a fatal crash, yielding the worst result "
        "(TP=6, FP=8, TN=1, Accuracy 39%).",
        bold_parts=["v4", "worst result"],
    )

    _para(
        doc,
        "v5 (PROMPT_J, reordered 8-stage) and v6 (PROMPT_K, 4-stage PERCEIVE/TRACK/ASSESS/DECIDE): "
        "v5 was aborted after 9/18 clips due to parse failures (partial accuracy ~33%). "
        "v6 re-balanced predictions (TP=5, FP=4, TN=5, FN=4, Accuracy 56%) but did not surpass v2.",
        bold_parts=["v5", "v6"],
    )

    _para(
        doc,
        "v7 (PROMPT_ORACLE_DEEP_V3, GT-injected): A forensic deep-audit prompt that was told the ground-truth "
        "outcome and asked to produce a 4-stage kinematic reconstruction (spatiotemporal anchoring, "
        "looming analysis, causal chain, counterfactuals). Not a blind prediction -- designed to generate "
        "rich explanatory training data with 100% correct verdicts. All 18 clips completed without error.",
        bold_parts=["v7"],
    )

    _para(
        doc,
        "Key finding: PROMPT_G (v2), with its simpler 6-step CoT, consistently outperformed all more "
        "complex variants. Adding reasoning steps introduced 'escape hatches' that allowed the model to "
        "rationalise away genuine danger signals. This insight directly motivates the v8 architecture.",
        bold_parts=["PROMPT_G (v2)", "Key finding"],
    )

    # -----------------------------------------------------------------------
    # 3. ACCURACY IN CONTEXT
    # -----------------------------------------------------------------------
    _heading(doc, "3. Accuracy in Context -- Comparison with Literature", 1)

    _para(
        doc,
        "The ~61% balanced accuracy achieved with v2 should be interpreted in the context of the "
        "fundamental difficulty of the task and the state of the art across model families.",
    )

    _bullet(
        doc,
        "Supervised vision-only models (trained on thousands of labeled clips): "
        "BADAS 1.0 (Nexar, 2025) achieves 86% AP / 88% AUC -- current SOTA on Nexar, trained on 40k "
        "proprietary videos with ego-centric collision filtering. STAGNet (2025) achieves new SOTA on "
        "DAD / DoTA / DADA using spatio-temporal graph + LSTM. Our own MViT baseline achieves ~90% "
        "accuracy on Nexar through supervised training.",
        bold_parts=["BADAS 1.0", "STAGNet", "MViT baseline"],
    )

    _bullet(
        doc,
        "Zero-shot / VLM-based approaches (no task-specific training): iFinder (2025) reports up to 39% "
        "gains in accident reasoning via structured semantic grounding for dashcam video. AccidentBench "
        "(2025) shows that even GPT-4o and Gemini 2.5 Pro score only 24-30% on harder multi-modal accident "
        "reasoning tasks. Our Gemini-3 Pro + PROMPT_G achieves ~61% on a balanced 18-clip set.",
        bold_parts=["iFinder", "AccidentBench", "Gemini-3 Pro + PROMPT_G"],
    )

    _para(
        doc,
        "The inherent difficulty stems from three compounding factors: (a) no task-specific training data "
        "seen by the model, (b) a text bottleneck -- visual observations must be serialised into language "
        "before reasoning, losing fine-grained spatial information, and (c) low temporal resolution "
        "(16 frames at 256px) compared to dedicated video models processing full-resolution streams. "
        "Importantly, our goal is not to compete on zero-shot accuracy -- the teacher MMLM is used "
        "offline to generate high-quality reasoning traces for student distillation, not for deployment.",
        bold_parts=["text bottleneck", "our goal is not to compete on zero-shot accuracy"],
    )

    # -----------------------------------------------------------------------
    # 3.1 SafeVL
    # -----------------------------------------------------------------------
    _heading(doc, "3.1  SafeVL -- A Key Related Work", 2)

    _para(
        doc,
        "A particularly relevant concurrent work is SafeVL, published by NVIDIA's Autonomous Vehicle "
        "Research Group in December 2025 (Ma, Cao, Ding et al., preprint). SafeVL was identified during "
        "the preparation of this report and is directly applicable to our research context.",
        bold_parts=["SafeVL", "NVIDIA's Autonomous Vehicle Research Group", "December 2025"],
    )

    _para(
        doc,
        "What SafeVL does: SafeVL is a VLM-based safety evaluator for autonomous driving that takes "
        "dashcam video as input, produces structured chain-of-thought reasoning traces, and outputs a "
        "binary safe/unsafe decision. It was explicitly evaluated on the Nexar real-world collision "
        "dataset -- the same dataset used in this thesis -- achieving 76% accuracy in the zero-shot "
        "setting, representing a 20% improvement over prior VLM baselines on the same data. The "
        "framework consists of two components: (1) a Road-Graph Counterfactual Data Generation Engine "
        "that synthesises diverse counterfactual unsafe scenarios to address the scarcity of real "
        "collision data, and (2) an Object-centric Visual Reasoning Framework that fuses these "
        "counterfactual scenarios with safe driving data for training.",
        bold_parts=["76% accuracy", "Nexar real-world collision dataset",
                    "Road-Graph Counterfactual Data Generation Engine",
                    "Object-centric Visual Reasoning Framework"],
    )

    _para(
        doc,
        "How SafeVL differs from our research: SafeVL is designed as an inference-time safety evaluator "
        "that requires a large VLM to be present at deployment. Our approach uses the VLM exclusively "
        "during offline teacher dataset construction, then distils its reasoning into a lightweight "
        "student model that operates without any LLM at inference time -- making it suitable for "
        "real-time ADAS deployment. Furthermore, SafeVL relies on a synthetic counterfactual data "
        "engine to produce training signal, whereas our Semi-Oracle Debate method generates correction "
        "signal organically from the model's own prediction errors on real Nexar clips, without any "
        "synthetic data pipeline.",
        bold_parts=["inference-time safety evaluator", "offline teacher dataset construction",
                    "lightweight student model", "organically from the model's own prediction errors"],
    )

    _para(
        doc,
        "How SafeVL benefits our research: SafeVL provides three concrete benefits. First, it validates "
        "that VLMs with structured CoT reasoning can perform meaningful collision prediction on Nexar, "
        "confirming that our thesis direction is both timely and technically sound. Second, the 15-point "
        "gap between our blind zero-shot result (61%) and SafeVL's 76% directly motivates the "
        "Semi-Oracle Debate approach (v8): if 76% is achievable with a purpose-built VLM framework, "
        "our debate-based correction mechanism -- which ensures every clip receives a correct verdict "
        "with plausible visual reasoning -- is a principled path toward closing this gap without "
        "synthetic data. Third, SafeVL's reliance on a large VLM at runtime underscores the exact "
        "deployment bottleneck that our student distillation pipeline is designed to solve, "
        "strengthening the practical motivation of this thesis.",
        bold_parts=["15-point gap", "Semi-Oracle Debate approach", "student distillation pipeline",
                    "deployment bottleneck"],
    )

    # -----------------------------------------------------------------------
    # 4. NOVELTY
    # -----------------------------------------------------------------------
    _heading(doc, "4. Novelty of Our Approach", 1)

    _para(
        doc,
        "The key novelty of this work is the Semi-Oracle with Debate method (v8) for teacher dataset "
        "construction. Rather than discarding clips where the blind prediction is wrong (losing ~40% of "
        "data at 61% accuracy), or injecting the ground truth directly (producing \"apologetic\" "
        "non-genuine reasoning), we use a debate framing to extract high-quality reasoning for all clips.",
        bold_parts=["Semi-Oracle with Debate method (v8)"],
    )

    _bullet(
        doc,
        "Not a zero-shot prediction system -- the goal is building a high-quality training corpus for "
        "student distillation. Prediction accuracy of the teacher at inference time is a diagnostic, "
        "not the end objective.",
        bold_parts=["Not a zero-shot prediction system"],
    )

    _bullet(
        doc,
        "Debate-based correction without ego-defense bias: when Pass 1 (blind PROMPT_G) disagrees with "
        "ground truth, Pass 2 sends a fresh API call asking the model to argue for the GT outcome in a "
        "formal debate framing -- without revealing it was wrong. This avoids the 'apologise and "
        "rationalise' pattern and instead produces the strongest possible evidence-based argument for "
        "the correct verdict.",
        bold_parts=["Debate-based correction without ego-defense bias"],
    )

    _bullet(
        doc,
        "Contrastive training pairs: mismatched clips produce both a wrong reasoning trace (Pass 1) and "
        "a correct debate-guided trace (Pass 2), creating natural contrastive examples that the student "
        "can learn from.",
        bold_parts=["Contrastive training pairs"],
    )

    _bullet(
        doc,
        "Uniform output schema: both passes use the identical PROMPT_G JSON structure, so the student "
        "model sees one consistent format across all training examples regardless of whether a clip "
        "came from a correct blind prediction or a debate correction.",
        bold_parts=["Uniform output schema"],
    )

    _bullet(
        doc,
        "100% coverage with grounded reasoning: every clip in the teacher dataset carries a correct "
        "verdict and a visually-grounded reasoning chain. This is unachievable with pure blind "
        "prediction (which discards errors) or direct oracle injection (which produces rationalised "
        "rather than genuinely derived reasoning).",
        bold_parts=["100% coverage with grounded reasoning"],
    )

    _para(
        doc,
        "Why this approach can achieve better downstream accuracy: the student model is trained not "
        "only on binary labels but on structured reasoning traces that identify specific visual cues "
        "(closing speed, lane transitions, occlusion, time-to-contact). Combined with a supervised "
        "vision backbone (CLIP encoder + temporal token mixer), the student learns WHAT to look for "
        "and WHY it predicts a collision, enabling it to generalise beyond the patterns memorised by "
        "a supervised classifier. The target accuracy range for the trained student is 85-92%, "
        "competitive with or exceeding the MViT baseline while adding explainability.",
        bold_parts=["85-92%", "WHAT to look for"],
    )

    # -----------------------------------------------------------------------
    # 5. CURRENT STATUS
    # -----------------------------------------------------------------------
    _heading(doc, "5. Current Status", 1)

    _para(
        doc,
        "The Semi-Oracle Debate v8 experiment is in progress. PROMPT_DEBATE_G has been added to "
        "prompts/templates.py and the script Teacher_dataset_distill_v8.py has been implemented. "
        "The 18-clip OpenRouter run is the immediate next step, which will produce "
        "teacher_dataset_v8.jsonl and teacher_dataset_v8.xlsx. Upon completion, Pass 1 accuracy "
        "will be compared against v2 (same PROMPT_G, confirming reproducibility) and the quality "
        "of debate-generated reasoning for mismatched clips will be reviewed manually.",
        bold_parts=["PROMPT_DEBATE_G", "Teacher_dataset_distill_v8.py",
                    "teacher_dataset_v8.jsonl", "teacher_dataset_v8.xlsx"],
    )

    # -----------------------------------------------------------------------
    # 6. FUTURE WORK
    # -----------------------------------------------------------------------
    _heading(doc, "6. Future Work", 1)

    _bullet(
        doc,
        "Scale the teacher dataset: run the Semi-Oracle Debate pipeline across the full Nexar dataset "
        "(hundreds of clips covering all TP time-to-event horizons and diverse TN scenarios) to build "
        "a corpus of sufficient size and diversity for student SFT.",
        bold_parts=["Scale the teacher dataset"],
    )

    _bullet(
        doc,
        "Human review loop: manually review and curate debate-generated reasoning to filter out cases "
        "where the model's arguments are not visually grounded, ensuring only high-quality reasoning "
        "traces enter the final training set.",
        bold_parts=["Human review loop"],
    )

    _bullet(
        doc,
        "Student model training: the existing MCAStudentModel architecture (CLIP vision encoder, "
        "temporal token mixer, causal LM projection, score head) will be trained via Supervised "
        "Fine-Tuning (SFT) on the curated teacher dataset. The loss combines binary cross-entropy "
        "on the collision score with language modelling loss on the reasoning text, teaching the "
        "student to both predict and explain.",
        bold_parts=["Student model training", "MCAStudentModel"],
    )

    _bullet(
        doc,
        "Fusion with MViT: the final deployed system may combine the MViT's high-accuracy binary "
        "prediction (~90%) with the student MMLM's reasoning capability in a confidence-weighted "
        "ensemble -- MViT provides precision, the student provides explainability and generalisation.",
        bold_parts=["Fusion with MViT"],
    )

    # -----------------------------------------------------------------------
    # Save
    # -----------------------------------------------------------------------
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Report saved to: {output_path}")


if __name__ == "__main__":
    here = Path(__file__).resolve().parent.parent  # MMLM_CursorAI/
    out = here / "outputs" / "progress_report.docx"
    build_report(out)
