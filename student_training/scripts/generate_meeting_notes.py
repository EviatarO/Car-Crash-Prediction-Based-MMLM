"""
generate_meeting_notes.py
=========================
Build a short DOCX of personal meeting notes that walk through the two
architecture diagrams (architecture_training.png + architecture_inference.png).
This is NOT a formal report — it is a quick-glance cheat-sheet for the
author to consult live during the meeting.

Output:
    outputs/reports/figures/E2_meeting_notes.docx

Sources of truth used to ground the answers:
    student_training/models/internvl_lora.py  (forward, get_score, ScoreHead)
    student_training/scripts/train_lora.py    (loss alpha, optimizer, scheduler)
    student_training/configs/train_lora.yaml  (lr, grad_accum, lora_r, lora_alpha)
    teacher_distillation/scripts/Teacher_dataset_distill_v11.py
"""

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT_ROOT = Path(__file__).resolve().parents[2]
FIG_DIR      = PROJECT_ROOT / "outputs" / "reports" / "figures"
OUT_DOCX     = FIG_DIR / "E2_meeting_notes.docx"

TRAIN_PNG = FIG_DIR / "architecture_training.png"
INFER_PNG = FIG_DIR / "architecture_inference.png"


# ── Helpers ──────────────────────────────────────────────────────────────────

def H1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)


def H2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)


def Q(doc, text):
    """Question line — bold dark-blue."""
    p = doc.add_paragraph()
    r = p.add_run("Q: " + text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x21, 0x4F, 0x86)


def A(doc, text):
    """Answer body — normal."""
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.size = Pt(10.5)


def B(doc, text):
    """Bullet."""
    p = doc.add_paragraph(text, style="List Bullet")
    for r in p.runs:
        r.font.size = Pt(10.5)


def CODE(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def IMG(doc, path: Path, width_in=6.5):
    if path.exists():
        doc.add_picture(str(path), width=Inches(width_in))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        A(doc, f"[missing image: {path.name}]")


# ── Build ────────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Tighter margins so the diagrams fit nicely
    for section in doc.sections:
        section.left_margin   = Inches(0.6)
        section.right_margin  = Inches(0.6)
        section.top_margin    = Inches(0.6)
        section.bottom_margin = Inches(0.6)

    H1(doc, "E2 LoRA — meeting cheat-sheet")
    A(doc, "Quick personal notes to glance at during the meeting. Not a report. "
           "Each answer is short; a bit longer where the concept is easy to mix up.")

    # ────────────────────────────────────────────────────────────────────────
    H1(doc, "Part 1 — architecture_training.png")
    IMG(doc, TRAIN_PNG, width_in=7.0)

    # 1. InternViT
    Q(doc, "1. What does InternViT do? Output? Sizes?")
    A(doc, "InternViT-300M is the frozen vision encoder (a Vision Transformer, ~300 M params). "
           "For each frame it tokenises the 448×448 image into a grid of 32×32 patches → "
           "1024 patch tokens, each a vector of dimension 1024 (the ViT hidden size).")
    B(doc, "Input  per frame:  3 × 448 × 448  RGB image (ImageNet-normalised).")
    B(doc, "Output per frame:  1024 patch tokens  ×  1024-dim each.")
    B(doc, "We feed 16 frames per clip, so the raw vision output is 16 × 1024 = 16,384 tokens "
           "before compression — way too many for the LLM context.")

    # 2. ViR pixel-shuffle
    Q(doc, "2. What does the ViR pixel-shuffle compression block do?")
    A(doc, "It shrinks the per-frame token count by a factor of 4 (a 2×2 spatial pixel-shuffle). "
           "It is a deterministic, parameter-free re-arrangement (no learned weights), so it is "
           "shown as FROZEN. Goal: make the visual sequence short enough to fit alongside the "
           "text prompt in the LLM context.")
    B(doc, "Input  per frame:  1024 patch tokens × 1024-dim.")
    B(doc, "Output per frame:  256 tokens × 4096-dim  (4× fewer tokens, 4× wider).")
    B(doc, "16 frames × 256 = 4096 visual tokens per clip — this is exactly the number "
           "of <IMG_CONTEXT> slots reserved in the prompt.")

    # 3. MLP projector
    Q(doc, "3. What does the MLP projector (mlp1) do?")
    A(doc, "A small 2-layer MLP that maps the visual embedding dimension (4096, after ViR) "
           "down to the LLM’s embedding dimension (2560 for Qwen3-4B). It’s the bridge that "
           "lets the language model 'read' image tokens as if they were normal text tokens.")
    B(doc, "Input :  4096 visual tokens × 4096-dim.")
    B(doc, "Output:  4096 visual tokens × 2560-dim  (= LLM hidden size).")
    B(doc, "It is TRAINABLE in our run (full-precision fine-tune of mlp1 in addition to LoRA).")

    # 4. LoRA params
    Q(doc, "4. LoRA parameters: q,k,v,o,gate,up,down, r=16, α=32 — what does each mean?")
    A(doc, "LoRA = Low-Rank Adaptation. Instead of fine-tuning a full weight matrix W "
           "(huge), we freeze W and add a tiny learnable update  ΔW = B · A,  where A is "
           "(r × in) and B is (out × r). With r=16, ΔW has only 2 · r · dim trainable params "
           "instead of dim · dim — that is the 1 % vs 100 % saving.")
    B(doc, "Target modules — the 7 linear layers inside every transformer block we attach LoRA to:")
    B(doc, "   • q, k, v, o   = Query / Key / Value / Output projections of the self-attention block.")
    B(doc, "   • gate, up, down = the three projections of the SwiGLU feed-forward (MLP) block.")
    B(doc, "   So we adapt BOTH attention AND the FFN — that is why the model can really learn new behaviour, not just re-weight attention.")
    B(doc, "r = 16  → rank of the low-rank update (capacity knob; bigger = more expressive, more params).")
    B(doc, "α = 32  → scaling factor; the effective update is  (α / r) · B·A = 2 · B·A. "
           "This decouples 'how much to update' from 'how big the matrices are'.")
    B(doc, "Net effect: only ~50 M trainable params (~1 % of the 4.94 B total) but applied "
           "in every transformer layer.")

    # 5. ScoreHead
    Q(doc, "5. How is the ScoreHead computed?")
    A(doc, "ScoreHead is one tiny linear layer — Linear(hidden_size → 1) — that turns a "
           "single hidden-state vector into one scalar logit. The σ (sigmoid) is applied "
           "OUTSIDE the head so we can use BCE-with-logits during training (numerically stable).")
    A(doc, "The vector it reads is the LLM’s last-layer hidden state at ONE specific position: "
           "the boundary token where the assistant is about to start writing its reply "
           "(asst_start_pos). At that point the model has already attended over all 4096 visual "
           "tokens + the full prompt, so the vector is a compact summary of 'everything I’ve seen'.")
    CODE(doc,
         "boundary_hidden = last_hidden[batch_idx, asst_start_pos, :]   # (B, H)\n"
         "score_logit     = score_head(boundary_hidden.float())          # (B,) raw logit\n"
         "P(collision)    = sigmoid(score_logit)                         # (B,) ∈ [0,1]")
    B(doc, "At inference (get_score) we use the LAST non-pad token instead of asst_start_pos "
           "— same idea: 'the position right after the model has read everything'.")

    # 6. LM head + CE
    Q(doc, "6. The second path — LM head + Cross-Entropy on reasoning. Last word or whole sentence?")
    A(doc, "The LM head is the standard next-token prediction head of Qwen3-4B (Linear "
           "hidden_size → vocab_size). It outputs a logit over the vocabulary at every position.")
    A(doc, "Cross-entropy is computed on EVERY token of the teacher’s JSON reasoning, not just "
           "the last word. Standard SFT trick: the `labels` tensor is set to -100 on the system "
           "prompt + user prompt + image tokens (so they are IGNORED in the loss), and equals the "
           "true token id only on the assistant-reply tokens. Hugging Face then averages CE over "
           "exactly those non-ignored positions.")
    B(doc, "Effect: the model is taught to reproduce the teacher’s full JSON ('scene_context', "
           "'final_reasoning', verdict, …) token-by-token — not just the YES/NO at the end.")
    B(doc, "Final training loss:   L  =  α · BCE(score)  +  (1 − α) · CE(reasoning),   α = 0.5.")
    B(doc, "So the score head and the reasoning head are trained jointly through ONE backward pass.")

    # 7. Hyperparameters
    Q(doc, "7. AdamW · cosine LR · lr=2e-4 · grad-accum 8 · bf16 + grad-checkpoint — where does each come in?")
    A(doc, "These are the standard fine-tuning knobs. A 1-line refresher of where each one sits "
           "in the training loop:")
    B(doc, "AdamW — the optimiser. Computes parameter updates from gradients. 'W' = decoupled "
           "weight decay (L2 regularisation done correctly). Applied only to the ~50 M trainable "
           "parameters (LoRA + mlp1 + ScoreHead).")
    B(doc, "Cosine LR schedule — the learning rate is not constant. It starts at lr=2e-4, warms "
           "up briefly, then decays smoothly along a cosine curve toward 0 over the 50 epochs. "
           "Helps the model fine-converge in the last epochs.")
    B(doc, "lr = 2e-4 — the peak learning rate. Typical for LoRA on a 4 B-param VLM.")
    B(doc, "grad-accum = 8 — micro-batch size is 1 clip (because each clip = 16 frames × 4096 "
           "tokens, GPU memory is tight). We do 8 backward passes BEFORE one optimiser step, so "
           "the effective batch size is 8 clips. Gradients are summed across the 8 micro-steps.")
    B(doc, "bf16 (bfloat16) — model weights & activations are stored in 16-bit floats, ~½ the "
           "memory of fp32 with the same dynamic range as fp32 (unlike fp16). The ScoreHead is "
           "kept in fp32 for numerical stability of BCE.")
    B(doc, "Gradient checkpointing — during the forward pass, intermediate activations are NOT "
           "kept in memory. They are recomputed on-the-fly during the backward pass. Trade: "
           "~30 % slower per step, but lets a 4 B model fit on a single 24–32 GB GPU.")

    A(doc, "Bird’s-eye loop:  for epoch in 50:  for clip in train (shuffled):  "
           "forward (vision → ViR → mlp1 → LLM → both heads) → compute α·BCE + (1-α)·CE → "
           "backward (every step, gradients accumulate) → AdamW step every 8 micro-batches → "
           "scheduler step → after the epoch, validate on the held-out 20 % and log F1/AP/loss "
           "to epoch_metrics.jsonl. Best checkpoint by val_F1 = epoch 27 / step 270.")

    # ────────────────────────────────────────────────────────────────────────
    doc.add_page_break()
    H1(doc, "Part 2 — architecture_inference.png")
    IMG(doc, INFER_PNG, width_in=7.0)

    # 1. ScoreHead at inference
    Q(doc, "1. How is the ScoreHead computed from the hidden state at inference?")
    A(doc, "Exactly the same arithmetic as in training, just without labels:")
    CODE(doc,
         "lm_out      = model(pixel_values, input_ids, attention_mask, output_hidden_states=True)\n"
         "last_hidden = lm_out.hidden_states[-1]              # (B, seq_len, 2560)\n"
         "last_pos    = attention_mask.sum(dim=1) - 1          # last non-pad position\n"
         "vec         = last_hidden[:, last_pos, :]            # (B, 2560)\n"
         "logit       = score_head(vec.float())                # (B,) one scalar per clip\n"
         "P(collision) = sigmoid(logit)                        # (B,) ∈ [0,1]")
    B(doc, "ONE forward pass through the full vision+LLM stack gives us the score.")
    B(doc, "A SECOND call (model.generate) produces the JSON reasoning text — same hidden states "
           "of the same model, just decoded autoregressively through the LM head.")

    # 2. Metrics meaning
    Q(doc, "2. Metrics — Recall, Precision, F1, AP. How to refer to each.")
    A(doc, "All four are computed from the (score, ground-truth) pairs of the 677 test clips.")
    B(doc, "Precision  =  TP / (TP + FP)   →  'Of the clips I called collision, how many really were?'  "
           "High precision = few false alarms.")
    B(doc, "Recall     =  TP / (TP + FN)   →  'Of the real collisions, how many did I catch?'  "
           "High recall = few missed accidents (this is the safety-critical one).")
    B(doc, "F1         =  2 · P · R / (P + R)   →  the harmonic mean of P and R, single number "
           "that balances false alarms vs misses. Reported AT a chosen threshold (we report "
           "F1@0.5 = 0.39 and F1* = 0.69 at the optimal threshold 0.195).")
    B(doc, "AP (Average Precision) — sweep the threshold from 1 → 0, plot the Precision–Recall "
           "curve, take the area under it. ONE number that summarises the entire P/R trade-off "
           "WITHOUT picking a threshold. This is our primary thesis metric  →  test AP = 0.636.")
    A(doc, "Rule of thumb when speaking: 'AP measures the ranking quality of my scores; "
           "F1/Precision/Recall measure the quality of a YES/NO decision at one specific threshold.'")

    # 3. Why P/R were not in training curves
    Q(doc, "3. Why couldn’t we plot Recall and Precision per epoch in the training curves?")
    A(doc, "Because the per-epoch logger (epoch_metrics.jsonl) only persisted aggregated scalars "
           "— train_f1, val_f1, train_ap, val_ap, val_loss — and NOT the underlying confusion "
           "matrix counts (TP/FP/FN/TN). And here is the gotcha: from F1 alone you CANNOT recover "
           "Precision and Recall, because many different (P, R) pairs give the same F1.")
    B(doc, "AP is also threshold-free, so it doesn’t implicitly carry P and R either.")
    B(doc, "Fix already pushed for FUTURE runs — train_lora.py now logs train_precision, "
           "val_precision, train_recall, val_recall, AND the raw TP/FP/FN/TN counts per epoch. "
           "So next training run will have full P/R curves out of the box; the existing E2 run "
           "just predates that patch.")
    B(doc, "For the TEST set we DO have full P/R because evaluate_metrics.py works on the raw "
           "(score, label) pairs of the 677 clips — it can compute anything from scratch.")

    # ────────────────────────────────────────────────────────────────────────
    doc.add_page_break()
    H2(doc, "One-line cheat answers (super-short, in case you blank out)")
    B(doc, "InternViT  →  turns each frame into 1024 visual tokens (frozen).")
    B(doc, "ViR pixel-shuffle  →  compresses 1024 → 256 tokens per frame (frozen, no params).")
    B(doc, "MLP projector  →  maps visual dim 4096 → LLM dim 2560 (trainable bridge).")
    B(doc, "LoRA r=16/α=32 on q,k,v,o,gate,up,down  →  tiny low-rank deltas in attention + FFN; ~50 M params.")
    B(doc, "ScoreHead  →  Linear(2560 → 1) on the boundary hidden state, then σ → P(collision).")
    B(doc, "CE loss on reasoning  →  per-token CE over the WHOLE assistant JSON, not just the verdict.")
    B(doc, "Combined loss  =  0.5 · BCE(score) + 0.5 · CE(reasoning).")
    B(doc, "AP  =  area under the Precision–Recall curve (threshold-free); we got 0.636 on test.")
    B(doc, "P, R  =  decisions at one threshold; recoverable from raw scores, not from F1 alone.")

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(f"Saved: {OUT_DOCX}")


if __name__ == "__main__":
    build()

