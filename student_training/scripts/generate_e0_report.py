"""
generate_e0_report.py
---------------------
Generates a ~2-page Word (.docx) performance summary report for the
E0 Zero-Shot Baseline experiment.

Output: outputs/reports/E0_zero_shot_performance_report.docx

Usage:
    python student_training/scripts/generate_e0_report.py
"""

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

# ── Paths ──────────────────────────────────────────────────────────────────────
PROJECT_ROOT = Path(__file__).resolve().parents[2]
METRICS_TEST  = PROJECT_ROOT / "outputs" / "metrics" / "zero_shot_test"
METRICS_TRAIN = PROJECT_ROOT / "outputs" / "metrics" / "zero_shot_train"
OUT_DIR       = PROJECT_ROOT / "outputs" / "reports"
OUT_FILE      = OUT_DIR / "E0_zero_shot_performance_report.docx"

FIGURES = {
    "cm_test":    METRICS_TEST  / "confusion_matrix.png",
    "cm_train":   METRICS_TRAIN / "confusion_matrix.png",
    "pr_test":    METRICS_TEST  / "pr_curve.png",
    "roc_test":   METRICS_TEST  / "roc_curve.png",
    "score_dist": METRICS_TEST  / "score_distribution.png",
    "group_ap":   METRICS_TEST  / "group_ap_bar.png",
    "pr_train":   METRICS_TRAIN / "pr_curve.png",
    "roc_train":  METRICS_TRAIN / "roc_curve.png",
}


# ── Helpers ────────────────────────────────────────────────────────────────────

def load_metrics(path: Path) -> dict:
    with open(path / "metrics.json", encoding="utf-8") as f:
        return json.load(f)


def set_margins(doc: Document, inches: float = 1.0):
    from docx.shared import Inches as In
    for section in doc.sections:
        section.top_margin    = In(inches)
        section.bottom_margin = In(inches)
        section.left_margin   = In(inches)
        section.right_margin  = In(inches)


def heading(doc: Document, text: str, level: int = 2):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    return p


def body(doc: Document, text: str, italic: bool = False, size: int = 11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    p.paragraph_format.space_after = Pt(4)
    return p


def caption(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    return p


def add_figure(doc: Document, img_path: Path, cap: str, width: float = 5.5):
    if not img_path.exists():
        body(doc, f"[Figure not found: {img_path.name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(width))
    caption(doc, cap)


def add_side_by_side(doc: Document,
                     left_img: Path, left_cap: str,
                     right_img: Path, right_cap: str,
                     width: float = 2.8):
    """Two figures in a borderless 2-cell table."""
    tbl = doc.add_table(rows=2, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Remove borders
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for border in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                el = OxmlElement(f"w:{border}")
                el.set(qn("w:val"), "none")
                tcBorders.append(el)
            tcPr.append(tcBorders)

    # Images row
    for col_idx, (img_path, _) in enumerate([(left_img, left_cap), (right_img, right_cap)]):
        cell = tbl.cell(0, col_idx)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        if img_path.exists():
            run.add_picture(str(img_path), width=Inches(width))
        else:
            p.add_run(f"[Missing: {img_path.name}]")

    # Captions row
    for col_idx, (_, cap_text) in enumerate([(left_img, left_cap), (right_img, right_cap)]):
        cell = tbl.cell(1, col_idx)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cap_text)
        run.italic = True
        run.font.size = Pt(9)

    doc.add_paragraph()  # spacing


def add_metric_table(doc: Document, m_test: dict, m_train: dict):
    """Comparison table of all metrics."""
    rows = [
        ("Average Precision (AP) ★",
         f"{m_test['ap']:.4f}",
         f"{m_train['ap']:.4f}"),
        ("AUC-ROC",
         f"{m_test['auc_roc']:.4f}",
         f"{m_train['auc_roc']:.4f}"),
        ("F1 (threshold = 0.5)",
         f"{m_test['f1']:.4f}",
         f"{m_train['f1']:.4f}"),
        ("Precision (threshold = 0.5)",
         f"{m_test['precision']:.4f}",
         f"{m_train['precision']:.4f}"),
        ("Recall (threshold = 0.5)",
         f"{m_test['recall']:.4f}",
         f"{m_train['recall']:.4f}"),
        ("Accuracy",
         f"{m_test['accuracy']:.4f}",
         f"{m_train['accuracy']:.4f}"),
        (f"Optimal F1 (thr = {m_test['optimal_threshold']:.2f})",
         f"{m_test['optimal_f1']:.4f}",
         f"{m_train['optimal_f1']:.4f}"),
        ("Mean score — Positive clips",
         f"{m_test['mean_score_pos']:.4f}",
         f"{m_train['mean_score_pos']:.4f}"),
        ("Mean score — Negative clips",
         f"{m_test['mean_score_neg']:.4f}",
         f"{m_train['mean_score_neg']:.4f}"),
    ]

    tbl = doc.add_table(rows=len(rows) + 1, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = tbl.rows[0].cells
    for i, txt in enumerate(["Metric", "Test Set (n=677)", "Train-100 (n=100)"]):
        p = hdr[i].paragraphs[0]
        run = p.add_run(txt)
        run.bold = True
        run.font.size = Pt(10)
        hdr[i]._tc.get_or_add_tcPr()

    # Data rows
    for r_idx, (metric, val_test, val_train) in enumerate(rows):
        row = tbl.rows[r_idx + 1].cells
        row[0].paragraphs[0].add_run(metric).font.size = Pt(10)
        row[1].paragraphs[0].add_run(val_test).font.size = Pt(10)
        row[2].paragraphs[0].add_run(val_train).font.size = Pt(10)
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()


def add_group_table(doc: Document, m_test: dict):
    group_metrics = m_test.get("group_metrics", {})
    if not group_metrics:
        body(doc, "(No per-group data available.)")
        return

    label_map = {"0": "0.5s before collision", "1": "1.0s before collision", "2": "1.5s before collision"}

    tbl = doc.add_table(rows=len(group_metrics) + 1, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = tbl.rows[0].cells
    for i, txt in enumerate(["Group", "AP", "n (clips)"]):
        run = hdr[i].paragraphs[0].add_run(txt)
        run.bold = True
        run.font.size = Pt(10)

    for r_idx, (grp, vals) in enumerate(sorted(group_metrics.items())):
        row = tbl.rows[r_idx + 1].cells
        row[0].paragraphs[0].add_run(label_map.get(str(grp), f"Group {grp}")).font.size = Pt(10)
        row[1].paragraphs[0].add_run(f"{vals['ap']:.4f}").font.size = Pt(10)
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[2].paragraphs[0].add_run(str(vals['n'])).font.size = Pt(10)
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()


def add_experiment_summary_table(doc: Document):
    rows = [
        ("E0", "InternVL3.5-4B-Flash (zero-shot)", "0.5397", "0.5613", "Baseline — no fine-tuning"),
        ("E1", "TBD", "—", "—", "LoRA fine-tuning (100 clips)"),
        ("E2", "TBD", "—", "—", "LoRA fine-tuning (200 clips)"),
        ("E3", "TBD", "—", "—", "LoRA fine-tuning (500 clips)"),
    ]

    tbl = doc.add_table(rows=len(rows) + 1, cols=5)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = tbl.rows[0].cells
    for i, txt in enumerate(["Exp.", "Model", "AP (Test)", "AUC (Test)", "Notes"]):
        run = hdr[i].paragraphs[0].add_run(txt)
        run.bold = True
        run.font.size = Pt(9)

    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            run = row[c_idx].paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            if c_idx in [2, 3]:
                row[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()


# ── Main report builder ────────────────────────────────────────────────────────

def build_report():
    m_test  = load_metrics(METRICS_TEST)
    m_train = load_metrics(METRICS_TRAIN)

    doc = Document()
    set_margins(doc, 1.0)

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_heading("E0 Baseline: Zero-Shot Performance Evaluation", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)

    sub = doc.add_paragraph("InternVL3.5-4B-Flash  |  Nexar Dashcam Dataset  |  April 2026")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].italic = True
    sub.paragraph_format.space_after = Pt(10)

    # ── Section 1: Evaluation Setup ───────────────────────────────────────────
    heading(doc, "1. Evaluation Setup")
    body(doc,
         "The E0 experiment evaluates InternVL3.5-4B-Flash in a fully zero-shot regime — "
         "no task-specific fine-tuning, no labeled examples. The model receives 16 frames "
         "(stride=4, bfloat16, resolution 448×448) and PROMPT_G, a 6-step Chain-of-Thought "
         "prompt, producing structured JSON output with a collision_verdict (YES/NO) and "
         "confidence (HIGH/MEDIUM/LOW). The verdict+confidence pair is mapped to a scalar "
         "score (YES/HIGH=0.90, YES/MED=0.65, YES/LOW=0.40, NO/HIGH=0.10, NO/MED=0.35, "
         "NO/LOW=0.60) used for AP and AUC computation.")
    body(doc,
         "Two evaluation splits are used: the Test Set (677 clips: 338 TP + 339 TN, from "
         "test.csv Usage='Private') and Train-100 (100 teacher-labelled clips: 50 TP + 50 TN, "
         "from manifest_v11). Clips are pre-cut by the Nexar dataset to end at 0.5s, 1.0s, "
         "or 1.5s before the collision event (or at matched non-collision midpoints). "
         "The primary thesis metric is Average Precision (AP); the target is AP > 0.67.")

    # ── Section 2: Core Metrics ───────────────────────────────────────────────
    heading(doc, "2. Core Performance Metrics")
    add_metric_table(doc, m_test, m_train)

    body(doc,
         "Metric definitions: "
         "(1) Average Precision (AP, ★ primary): area under the Precision-Recall curve across all "
         "thresholds, computed via sklearn.average_precision_score. Summarises performance without "
         "requiring a fixed threshold; higher is better; random baseline = 0.50 on balanced data. "
         "(2) AUC-ROC: area under the ROC curve (TPR vs FPR), measuring rank-order discrimination. "
         "(3) F1 at thr=0.5: harmonic mean of Precision and Recall at the standard 0.5 cut-off. "
         "(4) Optimal F1: best F1 achievable by sweeping thresholds on the PR curve — indicates the "
         "model's ceiling performance if threshold is tuned post-hoc. "
         "(5) Mean score (positive/negative): average raw output score per class; the gap between "
         "classes measures score separation. An ideal model would have mean_pos ≈ 0.9, mean_neg ≈ 0.1.")

    # ── Section 3: Confusion Matrices ────────────────────────────────────────
    heading(doc, "3. Confusion Matrices (threshold = 0.5)")
    add_side_by_side(
        doc,
        FIGURES["cm_test"],  "Figure 1: Test Set (n=677)",
        FIGURES["cm_train"], "Figure 2: Train-100 (n=100)",
        width=2.8,
    )
    body(doc,
         f"At the default threshold of 0.5, the model predicts negative (NO) the vast majority "
         f"of the time. On the test set, {m_test['fn']} of {m_test['n_positive']} true-positive clips are missed (FN), while "
         f"{m_test['tn']} of {m_test['n_negative']} negatives are correctly rejected (TN). TP recall is only "
         f"{m_test['recall']*100:.1f}%, while TN specificity is "
         f"{m_test['tn']/(m_test['tn']+m_test['fp'])*100:.1f}% ({m_test['tn']}/{m_test['n_negative']}). "
         f"This strong TN bias repeats consistently on Train-100 (TP recall = {m_train['recall']*100:.1f}%). "
         f"The pattern is not caused by class imbalance (both splits are balanced 50/50) but by "
         f"the model's intrinsic calibration — instruction-tuned VLMs tend to output "
         f"moderate, non-committal scores for uncertain inputs, pulling the score below 0.5 for "
         f"many real collision clips.")

    # ── Section 4: ROC + PR Curves ────────────────────────────────────────────
    heading(doc, "4. ROC and Precision-Recall Curves (Test Set)")
    add_side_by_side(
        doc,
        FIGURES["pr_test"],  "Figure 3: Precision-Recall Curve (AP=0.5397)",
        FIGURES["roc_test"], "Figure 4: ROC Curve (AUC=0.5613)",
        width=2.8,
    )
    body(doc,
         "Both curves are modestly above the random baseline (PR random = 0.50 horizontal line; "
         "ROC random = diagonal). The PR curve's sawtooth pattern reflects the coarse 4-level "
         "score vocabulary (0.10, 0.35, 0.65, 0.90) — a trained model with a continuous ScoreHead "
         "output will produce smoother curves and higher AP.")

    # ── PAGE BREAK ────────────────────────────────────────────────────────────
    doc.add_page_break()

    # ── Section 5: Score Distribution ─────────────────────────────────────────
    heading(doc, "5. Score Distribution by Class (Test Set)")
    add_figure(doc, FIGURES["score_dist"],
               "Figure 5: Model output score distribution — Collision (TP) vs Safe (TN) clips.",
               width=5.5)
    body(doc,
         f"The positive and negative score distributions substantially overlap. The mean score "
         f"for collision clips is {m_test['mean_score_pos']:.4f} vs "
         f"{m_test['mean_score_neg']:.4f} for safe clips — a separation of only "
         f"{m_test['mean_score_pos'] - m_test['mean_score_neg']:.4f}. "
         f"A well-calibrated model would show bimodal distributions clustered near 0 and 1. "
         f"The concentration of both classes at score = 0.35 (NO/MEDIUM verdict) indicates "
         f"the model defaults to low-confidence negative predictions for ambiguous clips. "
         f"The narrow score gap directly explains the depressed AP at this stage.")

    # ── Section 6: Per-Group Analysis ─────────────────────────────────────────
    heading(doc, "6. Per-Group Analysis — Time Before Collision (Test Set)")
    add_figure(doc, FIGURES["group_ap"],
               "Figure 6: AP by temporal group — clips closer to collision are easier to classify.",
               width=4.5)
    add_group_table(doc, m_test)
    body(doc,
         "AP decreases monotonically as the temporal distance from the collision increases: "
         "0.5s (AP=0.558) → 1.0s (AP=0.546) → 1.5s (AP=0.508). This is the expected direction: "
         "clips captured 0.5s before impact contain more unambiguous visual cues "
         "(vehicles in contact, occupant bracing, dramatic proximity violations) that a general "
         "vision-language model can partially detect. At 1.5s the visual scene is often "
         "indistinguishable from normal driving, making zero-shot detection near-random. "
         "The 0.05 AP spread across groups is modest at this stage. After fine-tuning, "
         "a steeper temporal gradient would confirm the model has learned collision-predictive "
         "features rather than scene-level confounds.")

    # ── Section 7: Reality Check & Literature Context ─────────────────────────
    heading(doc, "7. Reality Check and Literature Context")

    heading(doc, "7.1  Comparison to Supervised Baselines", level=3)
    body(doc,
         "State-of-the-art supervised models for dashcam collision anticipation (DSA, DSTA, "
         "UString) evaluated on standard benchmarks (DAD, CCD, A3D) typically report AP in "
         "the range 0.74–0.93. The E0 AP of 0.54 is substantially below these figures — but "
         "this comparison is not informative: those models are trained end-to-end on thousands "
         "of labeled collision clips, while E0 uses a 4B-parameter VLM with a text prompt only. "
         "The meaningful comparison is against zero-shot or few-shot VLM baselines. "
         "Published zero-shot VLM results on similar dashcam tasks report AP in the 0.50–0.60 "
         "range (limited benchmarks available). E0's AP of 0.54 is therefore consistent with "
         "the zero-shot VLM operating point and constitutes a credible starting baseline.")

    heading(doc, "7.2  TN Bias and Threshold Calibration", level=3)
    body(doc,
         "The optimal decision threshold is 0.35 on both splits (vs the default 0.5), "
         "recovering Optimal F1 ≈ 0.68. This should not be interpreted as deployable "
         "performance: the threshold was selected in-sample on the same data. "
         "The bias is a known property of RLHF-tuned instruction models — they are calibrated "
         "to be cautious and produce non-committal scores for uncertain inputs. "
         "Fine-tuning with binary BCE loss and a dedicated ScoreHead will re-calibrate "
         "the model's output distribution toward the collision detection task.")

    heading(doc, "7.3  Zero-Shot Limitations", level=3)
    body(doc,
         "InternVL3.5-4B-Flash has received no supervision signal for what constitutes 'imminent "
         "collision' vs 'normal driving'. The model likely relies on salient but imperfect proxies: "
         "road congestion, unusual camera angles, pedestrian presence — features correlated with "
         "but not directly causal of collision. Critically, the model has no access to motion "
         "cues or closing velocity; it reasons from static appearance differences across frames. "
         "These limitations directly motivate the fine-tuning experiments in E1–E3.")

    heading(doc, "7.4  Reliability Assessment", level=3)
    body(doc,
         "Results are internally consistent: both splits show AP 0.54–0.56, AUC 0.56–0.60, "
         "and optimal threshold 0.35, confirming the model is not overfit to any particular "
         "split. AP is meaningfully above random chance (0.50 on balanced data), confirming "
         "the model extracts some discriminative signal. The primary concern is the gap to "
         "the thesis target: current AP = 0.5397 vs target AP > 0.67 — a deficit of 0.13 "
         "that fine-tuning with teacher-distilled data is intended to close.")

    # ── Section 8: Experiment Summary ─────────────────────────────────────────
    heading(doc, "8. Experiments Summary (to be updated)")
    add_experiment_summary_table(doc)
    body(doc,
         "★ AP is the primary thesis metric. Target: AP > 0.67 on the held-out Test Set (n=677).",
         italic=True, size=9)

    # ── Save ──────────────────────────────────────────────────────────────────
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_FILE))
    print(f"Report saved: {OUT_FILE}")


if __name__ == "__main__":
    build_report()
